import docx
doc = docx.Document('/app/ESG_보고서_템플릿.docx')
with open('/app/tables_out.txt', 'w', encoding='utf-8') as f:
    f.write(f'Total tables: {len(doc.tables)}\n')
    for i, tbl in enumerate(doc.tables):
        try:
            if len(tbl.rows) > 0 and len(tbl.rows[0].cells) > 0:
                text = tbl.rows[0].cells[0].text.strip().replace('\n', ' ')
                f.write(f'Table {i}: {text[:80]}\n')
        except:
            pass
