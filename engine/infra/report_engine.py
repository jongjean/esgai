import os
import re
import json
import traceback
from datetime import datetime
from docx import Document
from docx.shared import Pt
from weasyprint import HTML
from jinja2 import Environment, FileSystemLoader

class ESGReportEngine:
    def __init__(self, template_dir: str):
        self.template_dir = template_dir
        self.docx_template_path = os.path.join(
            os.path.dirname(template_dir),
            "ESG_보고서_템플릿.docx"
        )
        self.html_template_name = "report.html"
        self.jinja_env = Environment(loader=FileSystemLoader(template_dir))
        print(f"[ENGINE] Initialized. Template Path: {self.docx_template_path}", flush=True)

    def _replace_text(self, container, search_str, replace_str):
        """
        문서 구조(w:sdt/컨텐츠 컨트롤)를 파괴하지 않기 위해 
        최대한 Run(w:r) 단위로만 최소한의 치환을 수행합니다.
        """
        if not container or not search_str:
            return

        # Paragraph 단위 처리
        paragraphs = container.paragraphs if hasattr(container, 'paragraphs') else [container]
        
        for p in paragraphs:
            # 1. 단일 Run 내에 검색어가 완벽히 포함된 경우 (가장 안전)
            for run in p.runs:
                if search_str in run.text:
                    run.text = run.text.replace(search_str, replace_str)
                    # 이미 치환되었다면 다음 Run으로 넘어감 (중복 방지)
            
            # 2. 검색어가 여러 Run에 걸쳐 분할된 경우 (보통 보호 모드에서는 드물지만 대응)
            if search_str in p.text:
                full_text = p.text
                new_text = full_text.replace(search_str, replace_str)
                # python-docx의 p.text = ... 는 런 서식을 모두 초기화할 수 있으므로
                # 최후의 수단으로만 사용하거나 루프 후에도 남아 있을 때만 제한적 적용
                # 여기서는 템플릿의 서식을 최대한 보호하기 위해 위 루프에서 해결되지 않은 경우만 대응
                if search_str in p.text:
                    p.text = new_text

    def generate_docx(self, data: dict, output_path: str):
        try:
            print(f"[ENGINE] Generating DOCX (Template-Preservation Mode): {output_path}", flush=True)
            company_name = data.get("company_name", "미등록 기업")
            industry     = data.get("industry", "미분류")
            raw_report   = data.get("raw_report", "")
            date_str     = datetime.now().strftime("%Y-%m-%d")
            size_label   = data.get("size_label", "미지정")

            # AI 생성 데이터 추출
            intro     = data.get("company_intro", "")
            products  = data.get("key_products", "")
            locations = data.get("locations", "")
            direction = data.get("esg_direction", "")

            # JSON 파싱
            try:
                payload = json.loads(raw_report) if isinstance(raw_report, str) else raw_report
                if "structured_data" in payload: payload = payload["structured_data"]
            except:
                payload = {}

            env = payload.get("environment", {})
            soc = payload.get("social", {})
            gov = payload.get("governance", {})

            # Fallback 데이터 구성
            intro     = intro or f"{company_name}은 {industry} 분야의 선도 기업입니다."
            products  = products or "핵심 제품 및 서비스"
            locations = locations or "전국 주요 사업장 및 운영 사이트"
            direction = direction or "지속가능한 성장을 위한 가치 창출"

            if not os.path.exists(self.docx_template_path):
                raise FileNotFoundError(f"Template not found at {self.docx_template_path}")

            # 템플릿 로드 (원본의 Protection 설정 유지)
            doc = Document(self.docx_template_path)
            
            # 치환 전략: 정규표현식(정보 블록) + 고정 플레이스홀더
            info_rules = [
                (re.compile(r'(기업\(기관\)명|회사\(기관\)명)\s*:?'), f"기업(기관)명 : {company_name}"),
                (re.compile(r'(기업\(기관\)형태|기업\(기관\)규모|회사\(기관\)분류)\s*:?'), f"기업(기관)형태 : {size_label}"),
                (re.compile(r'(산업분류|산업분야)\s*:?'), f"산업분류 : {industry}"),
                (re.compile(r'보고기간\s*:?'), f"보고기간: {datetime.now().year}년"),
                (re.compile(r'작성일\s*:?'), f"작성일: {date_str}"),
            ]
            
            anchor_mapping = {
                "ESG 경영 보고서 초안 템플릿": f"{company_name} ESG 경영 보고서",
                "[회사명]": company_name,
                "[업종]": industry,
                "[주요 제품/서비스]": products,
                "[주요 제품 또는 서비스]": products,
                "[사업장 또는 운영 현황]": locations,
                "[주요 지역]": locations,
                "[ESG 경영 방향]": direction,
                "[회사 소개]": intro,
                "[environment.activity]": str(env.get("activity", "친환경 공정 도입 및 에너지 효율화 추진")),
                "[environment.plan]": str(env.get("plan", "탄소 중립 달성을 위한 중장기 로드맵 이행")),
                "[environment.kpi]": str(env.get("kpi", "에너지 사용량 및 재생에너지 비중 관리")),
                "[social.policy]": str(soc.get("policy", "전 임직원 대상 공정 성과 체계 및 인권 경영 강화")),
                "[social.safety]": str(soc.get("safety", "사업장 정기 안전 점검 및 사고 제로화 실현")),
                "[social.kpi]": str(soc.get("kpi", "지역 사회 기여도 및 협력사 상생 협력 지수")),
                "[governance.system]": str(gov.get("system", "이사회 독립성 강화 및 책임 경영 체계 구축")),
                "[governance.ethics]": str(gov.get("ethics", "윤리 규범 준수 및 투명한 기업 문화 확산")),
            }

            def process_text_container(container):
                # 1. 고정 정보 블록 (Regex 기반 - 각 항목당 1회)
                for pattern, replacement in info_rules:
                    if pattern.search(container.text):
                        match = pattern.search(container.text)
                        original_tag = match.group(0)
                        self._replace_text(container, original_tag, replacement)

                # 2. 가변 플레이스홀더 (고정 문자열)
                for search_str, replace_str in anchor_mapping.items():
                    if search_str in container.text:
                        self._replace_text(container, search_str, replace_str)

            # 1. Paragraph 순회
            for para in doc.paragraphs:
                process_text_container(para)

            # 2. Table Cell 순회
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        process_text_container(cell)

            doc.save(output_path)
            print(f"[ENGINE] DOCX Saved Successfully (Preserved): {output_path}", flush=True)

        except Exception as e:
            print(f"❌ [ENGINE] DOCX ERROR: {e}", flush=True)
            traceback.print_exc()
            raise

    def generate_pdf(self, data: dict, output_path: str):
        try:
            print(f"[ENGINE] Generating PDF with Notice Box: {output_path}", flush=True)
            company_name = data.get("company_name", "미등록 기업")
            industry     = data.get("industry", "미분류")
            raw_report   = data.get("raw_report", "")
            date_str     = datetime.now().strftime("%Y-%m-%d %H:%M")
            size_label   = data.get("size_label", "미지정")

            try:
                payload = {}
                if raw_report:
                    try:
                        payload = json.loads(raw_report) if isinstance(raw_report, str) else raw_report
                        if "structured_data" in payload: payload = payload["structured_data"]
                    except:
                        payload = {}

                env = payload.get("environment") or data.get("environment", {})
                soc = payload.get("social") or data.get("social", {})
                gov = payload.get("governance") or data.get("governance", {})

                sections = []
                sections.append(f"<b>[기관 정보]</b><br>"
                               f"• 기업(기관)명 : {company_name}<br>"
                               f"• 기업(기관)형태 : {size_label}<br>"
                               f"• 산업분류 : {industry}")

                sections.append(f"<b>[환경 (Environment)]</b><br>"
                               f"• 실천 사항: {env.get('activity', '친환경 원자재 도입 및 에너지 효율화 실천')}<br>"
                               f"• 추진 계획: {env.get('plan', '탄소 중립 달성을 위한 로드맵 수립 및 실행')}<br>"
                               f"• 핵심 지표: {env.get('kpi', '온실가스 배출량 및 폐기물 재활용률 관리')}")
                
                sections.append(f"<b>[사회 (Social)]</b><br>"
                               f"• 인사 정책: {soc.get('policy', '임직원 복리후생 및 인권 정책 강화')}<br>"
                               f"• 안전 보건: {soc.get('safety', '사업장 안전 관리 시스템 구축 및 정기 점검')}<br>"
                               f"• 기여 지표: {soc.get('kpi', '지역 사회 공헌 지수 및 협력사 동반 성장')}")
                
                sections.append(f"<b>[거버넌스 (Governance)]</b><br>"
                               f"• 경영 체계: {gov.get('system', '이사회의 독립성 강화 및 책임 경영 체계 구축')}<br>"
                               f"• 윤리 보강: {gov.get('ethics', '윤리 규정 준수 및 반부패 문화 확산')}")
                
                pretty_report = "<br><br>".join(sections)
                pretty_report += "<br><br><small>※ 본 보고서는 KESGAI 엔진으로 자동 생성된 맞춤형 초안입니다.</small>"
            except Exception as e:
                pretty_report = str(raw_report) + f"<br><br>※ 데이터 처리 중 보고서 구조화 오류: {e}"

            context = {
                "company_name":    company_name,
                "industry":        industry,
                "generation_date": date_str,
                "full_report":     pretty_report
            }

            template = self.jinja_env.get_template(self.html_template_name)
            html_content = template.render(context)
            
            # PDF 생성
            temp_path = output_path + ".tmp"
            HTML(string=html_content).write_pdf(temp_path)
            os.replace(temp_path, output_path)
            print(f"[ENGINE] PDF Saved Successfully: {output_path}", flush=True)

        except Exception as e:
            print(f"❌ [ENGINE] PDF ERROR: {e}", flush=True)
            traceback.print_exc()
            raise
